#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Word文档上传下载测试脚本

此脚本用于：
1. 生成一个简单的Word文档
2. 通过SCP上传到服务器的summary_plans目录
3. 通过HTTP下载进行验证
"""

import os
import sys
import time
import paramiko
import requests
from scp import SCPClient
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_word_document():
    """
    创建一个测试用的Word文档
    """
    timestamp = time.strftime("%Y%m%d_%H%M%S")
    doc_filename = f"test_word_document_{timestamp}.docx"
    
    print(f"\n正在创建Word文档: {doc_filename}")
    
    # 创建文档对象
    doc = Document()
    
    # 添加标题
    title = doc.add_heading('SCP上传下载测试文档', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加正文段落
    doc.add_paragraph()  # 空行
    
    # 添加测试信息段落
    test_info = doc.add_paragraph()
    test_info_run = test_info.add_run('这是一个用于测试SCP上传和HTTP下载功能的Word文档。')
    test_info_run.font.bold = True
    
    # 添加详细信息
    doc.add_paragraph(f'创建时间: {timestamp}')
    doc.add_paragraph('测试流程:')
    doc.add_paragraph('1. 生成此Word文档')
    doc.add_paragraph('2. 通过SCP上传到服务器')
    doc.add_paragraph('3. 通过HTTP链接下载进行验证')
    
    # 添加表格
    doc.add_heading('测试配置信息', level=2)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    
    # 表头
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '配置项'
    hdr_cells[1].text = '值'
    
    # 表格内容
    rows = [
        ('服务器地址', '121.40.182.30'),
        ('SSH端口', '22'),
        ('用户名', 'batago'),
        ('目标目录', '/opt/redmine-3.0.1-0/apache2/htdocs/sharefile/summary_plans'),
        ('HTTP访问地址', 'http://121.40.182.30:8000/sharefile/summary_plans/'),
        ('测试时间', timestamp)
    ]
    
    for row_data in rows:
        row_cells = table.add_row().cells
        row_cells[0].text = row_data[0]
        row_cells[1].text = row_data[1]
    
    # 添加注意事项
    doc.add_heading('注意事项', level=2)
    doc.add_paragraph('此文档仅用于测试目的，用于验证SCP上传和HTTP下载功能是否正常工作。')
    doc.add_paragraph('如果您能成功下载并打开此文档，说明文件传输功能正常。')
    
    # 添加页脚
    sections = doc.sections
    for section in sections:
        footer = section.footer
        footer_paragraph = footer.add_paragraph()
        footer_run = footer_paragraph.add_run('SCP上传下载测试文档 - 自动生成')
        footer_run.font.size = Inches(8/72)  # 8pt 转换为英寸
        footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 保存文档
    doc.save(doc_filename)
    
    file_size = os.path.getsize(doc_filename)
    print(f"Word文档创建成功: {os.path.abspath(doc_filename)}")
    print(f"文件大小: {file_size} 字节 ({file_size/1024:.2f} KB)")
    
    return doc_filename

def connect_to_server():
    """
    连接到SSH服务器
    """
    try:
        # 创建SSH客户端
        ssh = paramiko.SSHClient()
        # 自动添加主机密钥
        ssh.set_missing_host_key_policy(paramiko.AutoAddPolicy())
        
        print("\n正在连接到SSH服务器...")
        # 连接服务器
        ssh.connect(
            hostname='121.40.182.30',
            port=22,
            username='batago',
            password='4008737505'
        )
        
        print("SSH连接成功!")
        return ssh
    except Exception as e:
        print(f"SSH连接失败: {str(e)}")
        return None

def create_remote_directory(ssh, directory):
    """
    在远程服务器上创建目录
    """
    try:
        print(f"\n正在创建远程目录: {directory}")
        # 使用shell命令创建目录（包括父目录）
        command = f"mkdir -p {directory}"
        stdin, stdout, stderr = ssh.exec_command(command)
        
        # 检查命令执行结果
        exit_code = stdout.channel.recv_exit_status()
        if exit_code == 0:
            print(f"远程目录创建成功: {directory}")
            return True
        else:
            error_msg = stderr.read().decode('utf-8')
            print(f"远程目录创建失败 (错误码: {exit_code}): {error_msg}")
            return False
    except Exception as e:
        print(f"创建远程目录时出错: {str(e)}")
        return False

def upload_file(ssh, local_file, remote_dir):
    """
    通过SCP上传文件
    """
    try:
        remote_filename = os.path.basename(local_file)
        remote_path = os.path.join(remote_dir, remote_filename).replace('\\', '/')
        
        print(f"\n正在通过SCP上传Word文档...")
        print(f"本地文件: {local_file}")
        print(f"远程路径: {remote_path}")
        
        # 开始计时
        start_time = time.time()
        
        # 创建SCP客户端
        with SCPClient(ssh.get_transport()) as scp:
            # 上传文件
            scp.put(local_file, remote_path)
        
        # 结束计时
        end_time = time.time()
        
        print(f"Word文档上传成功!")
        print(f"上传耗时: {end_time - start_time:.2f} 秒")
        return remote_path
    except Exception as e:
        print(f"Word文档上传失败: {str(e)}")
        return None

def verify_file_exists(ssh, remote_path):
    """
    验证文件是否存在于远程服务器上
    """
    try:
        print(f"\n验证远程文件是否存在: {remote_path}")
        # 执行ls命令检查文件
        command = f"ls -la {remote_path}"
        stdin, stdout, stderr = ssh.exec_command(command)
        
        # 检查命令执行结果
        exit_code = stdout.channel.recv_exit_status()
        if exit_code == 0:
            file_info = stdout.read().decode('utf-8')
            print(f"✓ 远程文件存在: {remote_path}")
            print(f"文件信息: {file_info.strip()}")
            return True
        else:
            error_msg = stderr.read().decode('utf-8')
            print(f"✗ 远程文件不存在: {remote_path}")
            print(f"错误信息: {error_msg}")
            return False
    except Exception as e:
        print(f"验证远程文件时出错: {str(e)}")
        return False

def test_http_download(remote_filename):
    """
    测试通过HTTP链接下载Word文档
    """
    # 构建HTTP下载链接
    http_link = f"http://121.40.182.30:8000/sharefile/summary_plans/{remote_filename}"
    
    print(f"\n正在测试HTTP下载Word文档...")
    print(f"HTTP下载链接: {http_link}")
    
    # 下载保存的本地文件名
    downloaded_filename = f"downloaded_{remote_filename}"
    
    try:
        # 发送HTTP GET请求
        print("发送HTTP请求...")
        start_time = time.time()
        response = requests.get(http_link, stream=True, timeout=60)
        end_time = time.time()
        
        # 检查响应状态
        if response.status_code == 200:
            # 保存下载的文件
            with open(downloaded_filename, 'wb') as f:
                for chunk in response.iter_content(chunk_size=8192):
                    if chunk:
                        f.write(chunk)
            
            download_time = end_time - start_time
            downloaded_size = os.path.getsize(downloaded_filename)
            
            print(f"✓ HTTP下载成功! 状态码: {response.status_code}")
            print(f"下载耗时: {download_time:.2f} 秒")
            print(f"下载的文件大小: {downloaded_size} 字节 ({downloaded_size/1024:.2f} KB)")
            print(f"文件保存为: {downloaded_filename}")
            
            # 验证文件大小
            original_size = os.path.getsize(remote_filename)
            print(f"原始文件大小: {original_size} 字节 ({original_size/1024:.2f} KB)")
            
            if abs(downloaded_size - original_size) < 1024:  # 允许1KB以内的差异
                print(f"✓ 文件大小验证通过! 差异: {abs(downloaded_size - original_size)} 字节")
                return True, http_link, downloaded_filename
            else:
                print(f"✗ 文件大小验证失败! 差异过大: {abs(downloaded_size - original_size)} 字节")
                return False, http_link, downloaded_filename
        else:
            print(f"✗ HTTP下载失败! 状态码: {response.status_code}")
            print(f"响应内容: {response.text[:200]}...")
            return False, http_link, None
    except Exception as e:
        print(f"HTTP下载请求出错: {str(e)}")
        return False, http_link, None

def main():
    """
    主函数
    """
    print("=== Word文档上传下载测试 ===")
    
    # 服务器信息
    server_info = {
        'hostname': '121.40.182.30',
        'port': 22,
        'username': 'batago',
        'password': '4008737505'
    }
    
    print("服务器信息:")
    print(f"- 地址: {server_info['hostname']}")
    print(f"- 端口: {server_info['port']}")
    print(f"- 用户名: {server_info['username']}")
    print(f"- 目标目录: /opt/redmine-3.0.1-0/apache2/htdocs/sharefile/summary_plans")
    print(f"- HTTP访问: http://121.40.182.30:8000/sharefile/summary_plans/")
    
    # 检查是否安装了python-docx
    try:
        import docx
        print("\npython-docx库检查: ✓ 已安装")
    except ImportError:
        print("\npython-docx库检查: ✗ 未安装")
        print("请先安装python-docx: pip install python-docx")
        return
    
    # 创建Word文档
    word_file = create_word_document()
    
    # SSH连接
    ssh = connect_to_server()
    if not ssh:
        print("\n测试失败: 无法连接到SSH服务器")
        if os.path.exists(word_file):
            os.remove(word_file)
            print(f"已清理本地Word文档: {word_file}")
        return
    
    # 远程目录
    remote_dir = "/opt/redmine-3.0.1-0/apache2/htdocs/sharefile/summary_plans"
    
    try:
        # 创建远程目录
        if not create_remote_directory(ssh, remote_dir):
            print("\n测试失败: 无法创建远程目录")
            return
        
        # 上传文件
        remote_path = upload_file(ssh, word_file, remote_dir)
        if not remote_path:
            print("\n测试失败: Word文档上传失败")
            return
        
        # 验证文件是否存在
        if not verify_file_exists(ssh, remote_path):
            print("\n测试失败: 远程文件验证失败")
            return
        
        # 测试HTTP下载
        remote_filename = os.path.basename(word_file)
        download_success, http_link, downloaded_file = test_http_download(remote_filename)
        
        if download_success:
            print("\n🎉 测试成功完成!")
            print("\n使用方法:")
            print(f"1. 通过HTTP直接下载: {http_link}")
            print(f"2. 通过SCP下载: scp batago@121.40.182.30:{remote_path} .")
            print(f"3. 本地已下载的文件: {downloaded_file}")
            print("\n请尝试打开下载的Word文档以确认内容完整性。")
        else:
            print("\n✗ HTTP下载测试失败")
            print("\n可能的原因:")
            print("1. Apache服务器未正确配置目录权限")
            print("2. 文件权限问题")
            print("3. 网络连接问题")
    
    finally:
        # 关闭SSH连接
        if ssh:
            ssh.close()
            print("\nSSH连接已关闭")
    
    print("\n=== 测试完成 ===")

if __name__ == "__main__":
    main()